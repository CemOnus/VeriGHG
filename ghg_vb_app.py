
# ghg_vb_app.py — Patched with sidebar "user_confirm" installer (warn-only)
import os
os.environ['PYARROW_IGNORE_DUPLICATE_FIELD_NAMES'] = '1'

import streamlit as st
import pandas as pd
import numpy as np
import io, re, math, json, datetime as dt, sys, subprocess

# ---- Dependency probe ----
try:
    import openpyxl  # Excel engine
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False

# Optional imports
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import requests
    from bs4 import BeautifulSoup
except Exception:
    requests = None
    BeautifulSoup = None

# PDF generation
REPORTLAB_AVAILABLE = True
try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
except Exception:
    REPORTLAB_AVAILABLE = False

st.set_page_config(page_title='GHG VB Prototype (Patched + Installer)', layout='wide')

EMISSION_FACTORS = {
    'diesel_liter': 2.68,
    'gasoline_liter': 2.31,
    'natural_gas_therm': 5.30,
    'electricity_kwh': 0.40,
    'air_travel_km': 0.15,
    'pgs_usd': 0.5,
    'commute_km': 0.14,
    '3pl_ton_km': 0.06
}

UNIT_CONVERSIONS = {
    'gallon_diesel': ('liter', 3.78541),
    'gallon_gasoline': ('liter', 3.78541),
    'm3_natgas': ('therm', 0.0366),
    'kwh': ('kwh', 1.0),
    'km': ('km', 1.0),
    'mile': ('km', 1.60934),
    'usd': ('usd', 1.0),
    'night': ('night', 1.0),
    'ton_km': ('ton_km', 1.0),
    'liter': ('liter', 1.0),
    'therm': ('therm', 1.0)
}

SCOPE_MAP = {
    'Fuel: Diesel': 'Scope 1',
    'Fuel: Gasoline': 'Scope 1',
    'Fuel: Natural Gas': 'Scope 1',
    'Electricity': 'Scope 2',
    'Business Travel': 'Scope 3',
    'Purchased Goods & Services': 'Scope 3',
    'Employee Commuting': 'Scope 3',
    'Third-Party Logistics': 'Scope 3'
}

# -----------------------------
# Session state
# -----------------------------
if 'records' not in st.session_state:
    st.session_state.records = []
if 'sources' not in st.session_state:
    st.session_state.sources = {}
if 'auditor_notes' not in st.session_state:
    st.session_state.auditor_notes = []
if 'quantified' not in st.session_state:
    st.session_state.quantified = None
if 'duplicate_log' not in st.session_state:
    st.session_state.duplicate_log = []

# -----------------------------
# Helpers
# -----------------------------
def to_float_safe(val):
    try:
        if isinstance(val, str):
            val = val.replace(',', '').strip()
        return float(val)
    except Exception:
        return None

def normalize_unit(value, unit_hint):
    if unit_hint is None:
        return None, None, None
    unit_hint = unit_hint.lower()
    if unit_hint in UNIT_CONVERSIONS:
        normalized, factor = UNIT_CONVERSIONS[unit_hint]
        return value * factor, normalized, factor
    for k, (norm, factor) in UNIT_CONVERSIONS.items():
        if k in unit_hint:
            return value * factor, norm, factor
    return value, unit_hint, 1.0

def _dedupe_list(names):
    seen = {}
    out = []
    for idx, n in enumerate(names):
        n0 = '' if n is None else str(n).strip()
        if not n0:
            n0 = f'column_{idx+1}'
        base = n0
        if base not in seen:
            seen[base] = 0
            out.append(base)
        else:
            seen[base] += 1
            out.append(f'{base}.{seen[base]}')
    return out

def dedupe_columns(df: pd.DataFrame, source_label: str = '') -> pd.DataFrame:
    if df is None or df.empty:
        return df
    orig_cols = [str(c) if c is not None else '' for c in df.columns]
    new_cols = _dedupe_list(orig_cols)
    if orig_cols != new_cols:
        st.session_state.duplicate_log.append({
            'source': source_label,
            'original': orig_cols,
            'renamed': new_cols
        })
    df = df.copy()
    df.columns = new_cols
    return df

def safe_show_df(df: pd.DataFrame, source_label: str = ''):
    if df is None:
        st.info('No data.')
        return df
    try:
        df_clean = dedupe_columns(df, source_label)
        st.dataframe(df_clean, use_container_width=True)
        return df_clean
    except Exception as e:
        st.warning(f'Could not render table for {source_label}: {e}')
        st.text(df.head(10).to_string(index=False) if len(df) else '(empty)')
        return df

# -----------------------------
# Parsing functions
# -----------------------------
def parse_excel(file_bytes, filename='(excel)'):
    dfs = {}
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes), engine=None)  # let pandas choose (openpyxl for .xlsx)
        for sheet in xls.sheet_names:
            try:
                df = xls.parse(sheet)
                dfs[sheet] = df
            except Exception as e:
                st.warning(f'Excel parse error in {filename}/{sheet}: {e}')
    except Exception as e:
        st.warning(f'Excel parse error in {filename}: {e}')
    return dfs

def parse_word(file_bytes):
    tables, text_blocks = [], []
    if docx is None:
        st.warning('python-docx not installed; cannot parse .docx.')
        return {'tables': tables, 'text': text_blocks}
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        for t in document.tables:
            rows = []
            for row in t.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                try:
                    df = pd.DataFrame(rows[1:], columns=rows[0])
                except Exception:
                    df = pd.DataFrame(rows)
                tables.append(df)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text_blocks = paragraphs
    except Exception as e:
        st.warning(f'Word parse error: {e}')
    return {'tables': tables, 'text': text_blocks}

def parse_pdf(file_bytes):
    tables, lines = [], []
    if pdfplumber is None:
        st.warning('pdfplumber not installed; cannot parse PDFs.')
        return {'tables': tables, 'lines': lines}
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    tbls = page.extract_tables() or []
                    for t in tbls:
                        try:
                            df = pd.DataFrame(t[1:], columns=t[0])
                        except Exception:
                            df = pd.DataFrame(t)
                        tables.append(df)
                except Exception:
                    pass
                try:
                    text = page.extract_text() or ''
                    for ln in text.splitlines():
                        if ln.strip():
                            lines.append(ln.strip())
                except Exception:
                    pass
    except Exception as e:
        st.warning(f'PDF parse error: {e}')
    return {'tables': tables, 'lines': lines}

def parse_url(url):
    if requests is None or BeautifulSoup is None:
        st.warning('requests/BeautifulSoup not installed; cannot fetch URLs.')
        return {'tables': [], 'text': ''}
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, 'html.parser')
        tables = []
        for html_table in soup.find_all('table'):
            rows = []
            for tr in html_table.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                if cells:
                    rows.append(cells)
            if rows:
                try:
                    df = pd.DataFrame(rows[1:], columns=rows[0])
                except Exception:
                    df = pd.DataFrame(rows)
                tables.append(df)
        text = soup.get_text(' ', strip=True)
        return {'tables': tables, 'text': text}
    except Exception as e:
        st.warning(f'URL parse error: {e}')
        return {'tables': [], 'text': ''}

# -----------------------------
# Extraction & quantification
# -----------------------------
def extract_emissions_records(from_df: pd.DataFrame, source_name: str):
    records = []
    if from_df is None or from_df.empty:
        return records
    df = from_df.copy()
    df.columns = [str(c).strip().lower() for c in df.columns]
    text_cols = [c for c in df.columns if any(k in c for k in ['desc','category','type','item','activity','fuel','name'])]
    qty_cols = [c for c in df.columns if any(k in c for k in ['qty','quantity','amount','value','usage','use','consumption'])]
    unit_cols = [c for c in df.columns if 'unit' in c or 'uom' in c]
    period_cols = [c for c in df.columns if any(k in c for k in ['period','month','date','year'])]
    note_cols = [c for c in df.columns if 'note' in c or 'comment' in c]
    if not text_cols and 'category' in df.columns:
        text_cols = ['category']
    if not qty_cols:
        for c in df.columns:
            if df[c].dtype.kind in 'fi':
                qty_cols.append(c); break
    for _, row in df.iterrows():
        text_blob = ' '.join([str(row[c]) for c in text_cols if c in df.columns]).lower()
        qty = None; unit_hint = None; period = None
        notes = ' '.join([str(row[c]) for c in note_cols if c in df.columns if not pd.isna(row[c])]).strip()
        for qc in qty_cols:
            val = row.get(qc, None)
            if pd.isna(val): 
                continue
            f = to_float_safe(val)
            if f is not None:
                qty = f; break
        for uc in unit_cols:
            u = row.get(uc, None)
            if not pd.isna(u):
                unit_hint = str(u)
        for pc in period_cols:
            p = row.get(pc, None)
            if not pd.isna(p):
                period = str(p)
        category = None
        if any(k in text_blob for k in ['diesel']):
            category = 'Fuel: Diesel'; unit_hint = unit_hint or 'liter'
        elif any(k in text_blob for k in ['gasoline','petrol']):
            category = 'Fuel: Gasoline'; unit_hint = unit_hint or 'liter'
        elif any(k in text_blob for k in ['natural gas','nat gas','ng']):
            category = 'Fuel: Natural Gas'; unit_hint = unit_hint or 'therm'
        elif any(k in text_blob for k in ['electric','kwh','grid']):
            category = 'Electricity'; unit_hint = unit_hint or 'kwh'
        elif any(k in text_blob for k in ['travel','flight','air','mileage']):
            category = 'Business Travel'; unit_hint = unit_hint or 'km'
        elif any(k in text_blob for k in ['purchased goods','pgs','spend','procure','supplier']):
            category = 'Purchased Goods & Services'; unit_hint = unit_hint or 'usd'
        elif any(k in text_blob for k in ['commute','employee commute']):
            category = 'Employee Commuting'; unit_hint = unit_hint or 'km'
        elif any(k in text_blob for k in ['logistics','3pl','freight','ton-km','ton_km']):
            category = 'Third-Party Logistics'; unit_hint = unit_hint or 'ton_km'
        if category and qty is not None:
            records.append({
                'source': source_name,
                'category': category,
                'activity_value': qty,
                'unit_hint': unit_hint,
                'period': period,
                'notes': notes
            })
    return records

def standardize_and_quantify(records):
    rows = []
    for r in records:
        cat = r['category']
        scope = SCOPE_MAP.get(cat, 'Scope 3')
        val = r['activity_value']
        unit_hint = (r['unit_hint'] or '').lower()
        norm_val, norm_unit, _ = normalize_unit(val, unit_hint)
        if cat == 'Fuel: Diesel':
            ef_key = 'diesel_liter'
        elif cat == 'Fuel: Gasoline':
            ef_key = 'gasoline_liter'
        elif cat == 'Fuel: Natural Gas':
            ef_key = 'natural_gas_therm'
        elif cat == 'Electricity':
            ef_key = 'electricity_kwh'
        elif cat == 'Business Travel':
            ef_key = 'air_travel_km'
        elif cat == 'Purchased Goods & Services':
            ef_key = 'pgs_usd'
        elif cat == 'Employee Commuting':
            ef_key = 'commute_km'
        else:
            ef_key = '3pl_ton_km'
        ef = EMISSION_FACTORS.get(ef_key, 0.0)
        kg_co2e = (norm_val or 0.0) * ef if norm_val is not None else None
        rows.append({
            'source': r['source'], 'category': cat, 'scope': scope,
            'activity_value': val, 'unit_hint': unit_hint,
            'normalized_value': norm_val, 'normalized_unit': norm_unit,
            'ef_used': ef_key, 'kg_co2e': kg_co2e,
            'period': r.get('period'), 'notes': r.get('notes','')
        })
    return pd.DataFrame(rows)

def verification_plan(df, assurance_level: str, materiality_pct: float):
    if df is None or df.empty:
        return pd.DataFrame(columns=['category','scope','kg_co2e','contribution_pct','risk','sample_size'])
    summary = df.groupby(['category','scope'], dropna=False)['kg_co2e'].sum().reset_index()
    total = summary['kg_co2e'].sum()
    summary['contribution_pct'] = np.where(total>0, 100.0*summary['kg_co2e']/total, 0.0)
    summary = summary.sort_values('contribution_pct', ascending=False)
    cum = 0.0; risks = []
    for _, row in summary.iterrows():
        cum += row['contribution_pct']
        if cum <= materiality_pct:
            risks.append('High')
        elif cum <= min(100.0, materiality_pct+20.0):
            risks.append('Moderate')
        else:
            risks.append('Low')
    summary['risk'] = risks
    base_n = 5 if assurance_level.lower().startswith('limited') else 15
    summary['sample_size'] = summary['risk'].map({'High': base_n+10, 'Moderate': base_n, 'Low': max(2, base_n//2)})
    return summary

def make_inventory_tables(df):
    scopes = df.groupby('scope')['kg_co2e'].sum().reset_index().sort_values('scope')
    cats = df.groupby(['scope','category'])['kg_co2e'].sum().reset_index()
    total = df['kg_co2e'].sum()
    return scopes, cats, total

def build_pdf(filename, title, sections):
    if not REPORTLAB_AVAILABLE:
        with open(filename.replace('.pdf','.txt'), 'w', encoding='utf-8') as f:
            f.write(title + '\\n' + '='*len(title) + '\\n\\n')
            for heading, body in sections:
                f.write(heading + '\\n' + '-'*len(heading) + '\\n')
                if isinstance(body, str):
                    f.write(body + '\\n\\n')
                else:
                    f.write(body.to_string(index=False) + '\\n\\n')
        return filename.replace('.pdf','.txt')
    doc = SimpleDocTemplate(filename, pagesize=LETTER, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet(); elems = [Paragraph(title, styles['Title']), Spacer(1,12)]
    for heading, body in sections:
        elems.append(Paragraph(heading, styles['Heading2'])); elems.append(Spacer(1,6))
        if isinstance(body, str):
            for para in body.split('\\n\\n'):
                elems.append(Paragraph(para, styles['BodyText'])); elems.append(Spacer(1,6))
        elif isinstance(body, pd.DataFrame) and not body.empty:
            data = [list(body.columns)] + body.fillna('').values.tolist()
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.25,colors.grey),('ALIGN',(0,0),(-1,-1),'LEFT'),('FONTSIZE',(0,0),(-1,-1),8)]))
            elems.append(t); elems.append(Spacer(1,12))
        else:
            elems.append(Paragraph('No data.', styles['BodyText'])); elems.append(Spacer(1,12))
    doc.build(elems); return filename

# -----------------------------
# Sidebar — Dependencies
# -----------------------------
with st.sidebar.expander("Dependencies", expanded=not OPENPYXL_AVAILABLE):
    if OPENPYXL_AVAILABLE:
        st.success("✅ Excel engine `openpyxl` is available.")
    else:
        st.warning("⚠️ Excel parsing requires `openpyxl`. Uploads remain enabled (warn-only), but `.xlsx` parsing will be skipped until installed.")
        if st.button("Install openpyxl now"):
            with st.spinner("Installing openpyxl..."):
                try:
                    result = subprocess.run([sys.executable, "-m", "pip", "install", "openpyxl"], capture_output=True, text=True, check=False)
                    if result.returncode == 0:
                        st.success("✅ Installed! Please **restart the app** and try again.")
                        st.code(result.stdout[:5000])
                    else:
                        st.error("❌ Installation failed. Your environment may not allow pip installs.")
                        st.code((result.stdout or '') + "\\n" + (result.stderr or ''))
                except Exception as e:
                    st.error(f"❌ Installation errored: {e}")

# -----------------------------
# UI - Engagement & Options
# -----------------------------
st.sidebar.title('GHG VB Prototype')
with st.sidebar.expander('Engagement & Assurance', expanded=True):
    engagement_type = st.selectbox('Type of Engagement', ['Verification','Validation'], index=0)
    assurance_level = st.selectbox('Level of Assurance', ['Limited','Reasonable'], index=0)
    materiality_pct = st.slider('Materiality threshold (% of total footprint)', 1, 20, 5)
with st.sidebar.expander('Options'):
    grid_ef = st.number_input('Grid EF (kg CO2e/kWh, for Electricity)', value=EMISSION_FACTORS['electricity_kwh'], step=0.01)
    EMISSION_FACTORS['electricity_kwh'] = grid_ef

# -----------------------------
# Main UI
# -----------------------------
st.title('GHG Verification Body – Prototype (Installer: warn-only)')
col_up1, col_up2 = st.columns(2)
with col_up1:
    uploaded = st.file_uploader('Upload files', type=['xlsx','xls','csv','docx','pdf'], accept_multiple_files=True)
with col_up2:
    url_input = st.text_input('Optional URL to parse:', placeholder='https://example.com/sustainability')

# Duplicate-log state exists from prior build; we keep it
def show_duplicate_log():
    st.header('Duplicate-Column Alerts (LOG Mode)')
    if st.session_state.duplicate_log:
        log_df = pd.DataFrame([{
            'source': e['source'],
            'original_cols': ', '.join(map(str, e['original']))[:300],
            'renamed_cols': ', '.join(map(str, e['renamed']))[:300]
        } for e in st.session_state.duplicate_log])
        st.dataframe(log_df, use_container_width=True, height=200)
        st.caption('Columns were auto-deduplicated (e.g., Amount, Amount.1, ...).')
    else:
        st.info('No duplicate or empty column headers detected yet.')

def extract_from_text_lines(lines, source_name):
    records = []
    patt = re.compile(r"(diesel|gasoline|natural gas|electricity|travel|commute|purchased goods|pgs|logistics|3pl|ton-?km)\s*[:\-]\s*([0-9\.,]+)\s*([A-Za-z_/-]+)?", re.I)
    for ln in lines:
        m = patt.search(ln)
        if m:
            kw = m.group(1).lower()
            qty = to_float_safe(m.group(2))
            unit = (m.group(3) or '').replace('/', '_').replace('-', '_')
            if 'diesel' in kw:
                cat, default = 'Fuel: Diesel', 'liter'
            elif 'gasoline' in kw:
                cat, default = 'Fuel: Gasoline', 'liter'
            elif 'natural gas' in kw:
                cat, default = 'Fuel: Natural Gas', 'therm'
            elif 'electricity' in kw:
                cat, default = 'Electricity', 'kwh'
            elif 'travel' in kw:
                cat, default = 'Business Travel', 'km'
            elif 'commute' in kw:
                cat, default = 'Employee Commuting', 'km'
            elif 'purchased goods' in kw or 'pgs' in kw:
                cat, default = 'Purchased Goods & Services', 'usd'
            else:
                cat, default = 'Third-Party Logistics', 'ton_km'
            if qty is not None:
                records.append({
                    'source': source_name,
                    'category': cat,
                    'activity_value': qty,
                    'unit_hint': unit or default,
                    'period': None,
                    'notes': ln
                })
    return records

if st.button('Parse Inputs'):
    for f in uploaded or []:
        source_name = f.name
        st.session_state.sources.setdefault(source_name, {'status':'Pending','notes':''})
        bytes_data = f.getvalue()
        fname = f.name.lower()

        if fname.endswith('.csv'):
            try:
                df = pd.read_csv(io.BytesIO(bytes_data))
                with st.expander(f'Parsed CSV: {f.name}'):
                    dfc = safe_show_df(df, f'{source_name}:CSV')
                    st.session_state.records.extend(extract_emissions_records(dfc, f'{source_name}:CSV'))
            except Exception as e:
                st.warning(f'CSV parse error for {f.name}: {e}')

        elif fname.endswith(('.xlsx','.xls')):
            if not OPENPYXL_AVAILABLE and fname.endswith('.xlsx'):
                st.warning(f"Skipping Excel .xlsx parsing for {f.name} because 'openpyxl' is not installed.")
            else:
                dfs = parse_excel(bytes_data, filename=f.name)
                with st.expander(f'Parsed Excel: {f.name}'):
                    for sname, df in dfs.items():
                        st.write(f'**Sheet:** {sname}')
                        dfc = safe_show_df(df, f'{source_name}:{sname}')
                        st.session_state.records.extend(extract_emissions_records(dfc, f'{source_name}:{sname}'))

        elif fname.endswith('.docx'):
            parsed = parse_word(bytes_data)
            with st.expander(f'Parsed Word: {f.name}'):
                for i, tdf in enumerate(parsed['tables']):
                    st.write(f'**Table {i+1}**')
                    tdfc = safe_show_df(tdf, f'{source_name}:table{i+1}')
                    st.session_state.records.extend(extract_emissions_records(tdfc, f'{source_name}:table{i+1}'))
                if parsed['text']:
                    st.write('**Paragraphs (first 10):**'); st.write(parsed['text'][:10])

        elif fname.endswith('.pdf'):
            parsed = parse_pdf(bytes_data)
            with st.expander(f'Parsed PDF: {f.name}'):
                for i, tdf in enumerate(parsed['tables']):
                    st.write(f'**Table {i+1}**')
                    tdfc = safe_show_df(tdf, f'{source_name}:table{i+1}')
                    st.session_state.records.extend(extract_emissions_records(tdfc, f'{source_name}:table{i+1}'))
                if parsed['lines']:
                    st.write('**Detected Lines (first 15):**'); st.write(parsed['lines'][:15])
                    st.session_state.records.extend(extract_from_text_lines(parsed['lines'], source_name))
        else:
            st.warning(f'Unsupported file type for {f.name}')

    if url_input.strip():
        parsed = parse_url(url_input.strip())
        with st.expander('Parsed URL'):
            for i, tdf in enumerate(parsed['tables']):
                st.write(f'**Table {i+1}**')
                tdfc = safe_show_df(tdf, f'URL:{url_input}:table{i+1}')
                st.session_state.records.extend(extract_emissions_records(tdfc, f'URL:{url_input}:table{i+1}'))
            if parsed['text']:
                st.write('**Extracted text sample (first 500 chars):**'); st.write(parsed['text'][:500])
    st.success('Parsing complete.')

# Show duplicate-column log
show_duplicate_log()

# Records & approvals
st.header('1) Extracted Activity Records')
if st.session_state.records:
    st.dataframe(pd.DataFrame(st.session_state.records), use_container_width=True, height=300)
else:
    st.info('No records yet. Upload and Parse.')

st.subheader('Source Approvals')
for source_name, meta in st.session_state.sources.items():
    with st.expander(source_name):
        c1, c2 = st.columns([1,2])
        with c1:
            status = st.selectbox('Status', ['Pending','Approved','Rejected'], index=['Pending','Approved','Rejected'].index(meta['status']), key=f'status_{source_name}')
        with c2:
            notes = st.text_area('Reviewer notes', value=meta.get('notes',''), key=f'notes_{source_name}')
        st.session_state.sources[source_name]['status'] = status
        st.session_state.sources[source_name]['notes'] = notes

# Quantification
st.header('2) Standardize & Quantify')
if st.button('Run Quantification'):
    df_records = pd.DataFrame(st.session_state.records)
    if df_records.empty:
        st.warning('No records to quantify.')
    else:
        st.session_state.quantified = standardize_and_quantify(st.session_state.records)
        st.success('Quantification complete.')

# Verification and reports
if st.session_state.quantified is not None:
    quantified = st.session_state.quantified
    st.dataframe(quantified, use_container_width=True, height=320)
    scopes, cats, total = make_inventory_tables(quantified)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown('**By Scope (kg CO2e)**'); st.dataframe(scopes, use_container_width=True)
    with c2:
        st.markdown('**By Category (kg CO2e)**'); st.dataframe(cats, use_container_width=True)
    with c3:
        st.metric('Total Footprint (t CO2e)', f'{total/1000.0:,.2f}')
    st.header('3) Verification Plan')
    vplan = verification_plan(quantified, assurance_level, materiality_pct)
    st.dataframe(vplan, use_container_width=True)
    st.caption('Risk bands and sample sizes are demonstrative only.')
    st.header('4) Auditor Annotations & Evidence')
    note = st.text_area('Add a finding / observation / assumption:')
    if st.button('Add Note'):
        if note.strip():
            st.session_state.auditor_notes.append({'ts': dt.datetime.now().isoformat(timespec='seconds'), 'note': note.strip()})
            st.success('Note added.')
    if st.session_state.auditor_notes:
        st.dataframe(pd.DataFrame(st.session_state.auditor_notes), use_container_width=True, height=200)
    evidence = st.file_uploader('Optional: Upload supporting evidence files (any type)', accept_multiple_files=True)
    if evidence:
        st.info('Evidence received: ' + ', '.join([f.name for f in evidence]))
    st.header('5) Generate Reports')
    org_name = st.text_input('Reporting Entity Name', value='Demo Company LLC')
    reporting_year = st.text_input('Reporting Period', value='FY 2024 (01-Jul-2024 to 30-Jun-2025)')
    verifier_name = st.text_input('Verification Body', value='Demo Verification Body, Inc.')
    materiality_str = f'{materiality_pct}% of total inventory'
    if st.button('Build PDF Reports'):
        ts = dt.datetime.now().strftime('%Y%m%d_%H%M%S')
        inv_sections = [
            ('Organizational Boundaries', f'Entity: {org_name}\\nOwnership/Control: Operational Control\\nReporting Period: {reporting_year}'),
            ('Operational Boundaries', 'Scope 1 (fuels), Scope 2 (electricity), Scope 3 (BT, PGS, commuting, 3PL).'),
            ('GHG Quantification', f'Illustrative EFs. Materiality: {materiality_str}. Assurance: {assurance_level}. Engagement: {engagement_type}'),
            ('By Scope (kg CO2e)', scopes),
            ('By Category (kg CO2e)', cats),
            ('Total', pd.DataFrame([{'Total kg CO2e': total, 'Total t CO2e': total/1000.0}])),
            ('Source Approvals', pd.DataFrame([{'source': s, **meta} for s, meta in st.session_state.sources.items()]))
        ]
        def build_pdf(filename, title, sections):
            if not REPORTLAB_AVAILABLE:
                with open(filename.replace('.pdf','.txt'), 'w', encoding='utf-8') as f:
                    f.write(title + '\\n' + '='*len(title) + '\\n\\n')
                    for heading, body in sections:
                        f.write(heading + '\\n' + '-'*len(heading) + '\\n')
                        if isinstance(body, str):
                            f.write(body + '\\n\\n')
                        else:
                            f.write(body.to_string(index=False) + '\\n\\n')
                return filename.replace('.pdf','.txt')
            doc = SimpleDocTemplate(filename, pagesize=LETTER, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
            styles = getSampleStyleSheet(); elems = [Paragraph(title, styles['Title']), Spacer(1,12)]
            for heading, body in sections:
                elems.append(Paragraph(heading, styles['Heading2'])); elems.append(Spacer(1,6))
                if isinstance(body, str):
                    for para in body.split('\\n\\n'):
                        elems.append(Paragraph(para, styles['BodyText'])); elems.append(Spacer(1,6))
                elif isinstance(body, pd.DataFrame) and not body.empty:
                    data = [list(body.columns)] + body.fillna('').values.tolist()
                    t = Table(data, repeatRows=1)
                    t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.25,colors.grey),('ALIGN',(0,0),(-1,-1),'LEFT'),('FONTSIZE',(0,0),(-1,-1),8)]))
                    elems.append(t); elems.append(Spacer(1,12))
                else:
                    elems.append(Paragraph('No data.', styles['BodyText'])); elems.append(Spacer(1,12))
            doc.build(elems); return filename
        inv_path = f'GHG_Inventory_{ts}.pdf'
        inv_file = build_pdf(inv_path, f'{org_name} – GHG Inventory Report', inv_sections)
        vstmt_sections = [
            ('Verification Statement', f'Provided by {verifier_name} to {org_name} for {reporting_year}. '
             f'Engagement: {engagement_type}, Assurance: {assurance_level}, Materiality: {materiality_str}. '
             'Based on procedures and evidence obtained, no material misstatement noted (demo language).'),
            ('Scope & Boundary', 'Scope 1, Scope 2, selected Scope 3 categories included.'),
            ('Summary of Results', pd.DataFrame([{'Total t CO2e': total/1000.0, 'Assurance': assurance_level, 'Materiality %': materiality_pct}]))
        ]
        vstmt_path = f'Verification_Statement_{ts}.pdf'
        vstmt_file = build_pdf(vstmt_path, f'{org_name} – Verification Statement', vstmt_sections)
        svrep_sections = [
            ('Objective, Criteria, and Scope', f'Objective: Verify {org_name} GHG assertion. Criteria: ISO 14064 & GHG Protocol. Scope: per boundaries.'),
            ('Methodology', 'Planning, risk assessment, sampling by contribution, analytical procedures, inspection, recalculation.'),
            ('Findings', 'No material inconsistencies noted in demo. Some unit assumptions applied.'),
            ('Verification Plan Snapshot', vplan),
            ('Auditor Notes (excerpt)', pd.DataFrame(st.session_state.auditor_notes) if st.session_state.auditor_notes else pd.DataFrame([{'note':'(none recorded in demo)'}]))
        ]
        svrep_path = f'Summary_Verification_Report_{ts}.pdf'
        svrep_file = build_pdf(svrep_path, f'{org_name} – Summary Verification Report', svrep_sections)
        st.success('Reports generated.')
        if inv_file.endswith('.pdf'):
            with open(inv_file, 'rb') as f: st.download_button('Download GHG Inventory Report (PDF)', f, file_name=inv_file, mime='application/pdf')
        else:
            with open(inv_file, 'r', encoding='utf-8') as f: st.download_button('Download GHG Inventory Report (TXT)', f, file_name=inv_file, mime='text/plain')
        if vstmt_file.endswith('.pdf'):
            with open(vstmt_file, 'rb') as f: st.download_button('Download Verification Statement (PDF)', f, file_name=vstmt_file, mime='application/pdf')
        else:
            with open(vstmt_file, 'r', encoding='utf-8') as f: st.download_button('Download Verification Statement (TXT)', f, file_name=vstmt_file, mime='text/plain')
        if svrep_file.endswith('.pdf'):
            with open(svrep_file, 'rb') as f: st.download_button('Download Summary Verification Report (PDF)', f, file_name=svrep_file, mime='application/pdf')
        else:
            with open(svrep_file, 'r', encoding='utf-8') as f: st.download_button('Download Summary Verification Report (TXT)', f, file_name=svrep_file, mime='text/plain')

st.caption('Patched build with sidebar dependency installer (warn-only), duplicate-column handling, and demo factors.')
